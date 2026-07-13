#!/usr/bin/env python3
"""
06_appendix_lexicon.py
======================
Turn the interactional-lexicon *audit* into a publication-ready appendix table:
the full term list by layer and source, with ambiguity flags and the terms
excluded after concordance inspection (with reasons).

Input  (newest, unless --audit given):
    ConnversationalRSs/outputs/interactional_lexicon_audit_<date>.csv
        columns: layer, term, source, ambiguous, excluded_after_inspection,
                 exclusion_reason, n_records, n_records_near_node

Outputs (date-stamped, into OUT_DIR unless --outdir given):
    appendix_interactional_lexicon_<date>.csv    tidy, human-readable columns
    appendix_interactional_lexicon_<date>.md      Markdown table (always)
    appendix_interactional_lexicon_<date>.docx    formatted Word table (if
                                                  python-docx is installed)

Notes
-----
* Per-TERM counts (n_records, n_records_near_node) are reported as-is; these are
  distinct-record counts for a single term and involve no de-duplication issue.
* The per-LAYER summary reports terms / excluded / attested only. It does NOT
  sum n_records into a record total: summing double-counts any record that uses
  two terms from the same layer. The record-level "records with >=1 term" figure
  (the gradient) must come from a union computed over record ids in the main
  lexical script, not from this table. See the note printed in the appendix.
* Deterministic: no randomness anywhere, so no seed is required.
"""
from __future__ import annotations
import sys, argparse
from datetime import date
from pathlib import Path
import pandas as pd

sys.path.insert(0, str(Path(__file__).resolve().parent))
try:
    from runlog import log_run
except Exception:                       # keep the script usable off-pipeline
    log_run = None

# ---------------------------------------------------------------- CONFIG -----
BASE_DIR = Path("ConnversationalRSs")
OUT_DIR  = BASE_DIR / "outputs"
RUN_DATE = date.today().isoformat()

# layer order = distance from computer science (matches the gradient figure)
LAYER_ORDER = ["CA_core", "pragmatics", "hci_cui"]
LAYER_NICE = {
    "CA_core":    "Conversation analysis (interaction-as-object)",
    "pragmatics": "Pragmatics (meaning-in-interaction)",
    "hci_cui":    "HCI / CUI (interaction-as-interface)",
}


def to_bool(s: pd.Series) -> pd.Series:
    """Coerce True/False that may have been read back as strings."""
    if s.dtype == bool:
        return s
    return s.astype(str).str.strip().str.lower().isin({"true", "1", "yes"})


def find_audit(explicit: str | None) -> Path:
    if explicit:
        p = Path(explicit)
        if not p.exists():
            sys.exit(f"ERROR: audit file not found: {p}")
        return p
    cands = sorted(OUT_DIR.glob("interactional_lexicon_audit_*.csv"))
    if not cands:
        sys.exit(f"ERROR: no interactional_lexicon_audit_*.csv in {OUT_DIR}")
    return cands[-1]                     # newest by date-stamped name


def status_label(row) -> str:
    if row["excluded_after_inspection"]:
        reason = str(row["exclusion_reason"]).strip() or "excluded after inspection"
        return f"Excluded — {reason}"
    if row["ambiguous"]:
        return "Validated (ambiguous; concordance-checked)"
    return "Validated"


def tidy_frame(audit: pd.DataFrame) -> pd.DataFrame:
    """Ordered, human-readable long table for CSV/MD/DOCX."""
    a = audit.copy()
    a["excluded_after_inspection"] = to_bool(a["excluded_after_inspection"])
    a["ambiguous"] = to_bool(a["ambiguous"])
    a["_lord"] = a["layer"].map({l: i for i, l in enumerate(LAYER_ORDER)}).fillna(99)
    # within a layer: attested first (desc), then zero-count terms alphabetically
    a = a.sort_values(["_lord", "n_records", "term"], ascending=[True, False, True])
    a["Status"] = a.apply(status_label, axis=1)
    out = a[["layer", "term", "source", "n_records", "n_records_near_node", "Status"]].copy()
    out.columns = ["Layer", "Term", "Source", "Records", "Records near a node word", "Status"]
    out["Layer"] = out["Layer"].map(LAYER_NICE).fillna(out["Layer"])
    return out


def layer_summary(audit: pd.DataFrame) -> pd.DataFrame:
    a = audit.copy()
    a["excluded_after_inspection"] = to_bool(a["excluded_after_inspection"])
    rows = []
    for layer in LAYER_ORDER:
        g = a[a["layer"] == layer]
        if g.empty:
            continue
        val = g[~g["excluded_after_inspection"]]
        rows.append({
            "Layer": LAYER_NICE[layer],
            "Source": g["source"].iloc[0],
            "Terms searched": len(g),
            "Excluded after inspection": int(g["excluded_after_inspection"].sum()),
            "Validated terms": len(val),
            "Validated terms attested (>=1 record)": int((val["n_records"] > 0).sum()),
        })
    return pd.DataFrame(rows)


CAPTION = (
    "Appendix X. Interactional lexicon: terms, sources, and concordance-based "
    "exclusions. Each term was searched over the title and abstract of all 866 "
    "records. Layers are theory-derived and kept disjoint (terms shared by "
    "conversation analysis and the interface literature are assigned to the "
    "conversation-analytic layer). 'Records' is the number of records in which "
    "the term occurs; 'Records near a node word' is the subset in which it "
    "co-occurs with a conversation-family node word within the concordance "
    "window. Ambiguous terms were checked against their concordances and, where "
    "not used in the interactional sense, excluded (reason given). Per-layer "
    "record totals are not the sum of the term rows (a record may contain "
    "several terms from one layer); the record-level counts underlying the "
    "gradient figure are reported in Section 6."
)


def write_markdown(path: Path, summary: pd.DataFrame, tidy: pd.DataFrame) -> None:
    lines = [f"# {CAPTION}", "", "## Summary by layer", "",
             summary.to_markdown(index=False), "", "## Full term list", ""]
    for layer in LAYER_ORDER:
        nice = LAYER_NICE[layer]
        sub = tidy[tidy["Layer"] == nice]
        if sub.empty:
            continue
        src = sub["Source"].iloc[0]
        lines += [f"### {nice}", f"*Source: {src}*", "",
                  sub[["Term", "Records", "Records near a node word", "Status"]]
                  .to_markdown(index=False), ""]
    path.write_text("\n".join(lines), encoding="utf-8")


def write_docx(path: Path, summary: pd.DataFrame, tidy: pd.DataFrame) -> str | None:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return None
    doc = Document()
    h = doc.add_heading("Appendix X. Interactional lexicon", level=1)
    cap = doc.add_paragraph(CAPTION)
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)

    doc.add_heading("Summary by layer", level=2)
    st = doc.add_table(rows=1, cols=len(summary.columns))
    st.style = "Light Grid Accent 1"
    for j, c in enumerate(summary.columns):
        st.rows[0].cells[j].paragraphs[0].add_run(str(c)).bold = True
    for _, r in summary.iterrows():
        cells = st.add_row().cells
        for j, c in enumerate(summary.columns):
            cells[j].text = str(r[c])

    doc.add_heading("Full term list", level=2)
    cols = ["Term", "Records", "Records near a node word", "Status"]
    for layer in LAYER_ORDER:
        nice = LAYER_NICE[layer]
        sub = tidy[tidy["Layer"] == nice]
        if sub.empty:
            continue
        src = sub["Source"].iloc[0]
        doc.add_heading(nice, level=3)
        p = doc.add_paragraph(f"Source: {src}")
        p.runs[0].italic = True
        t = doc.add_table(rows=1, cols=len(cols))
        t.style = "Light Grid Accent 1"
        for j, c in enumerate(cols):
            t.rows[0].cells[j].paragraphs[0].add_run(c).bold = True
        for _, r in sub.iterrows():
            cells = t.add_row().cells
            for j, c in enumerate(cols):
                cells[j].text = str(r[c])
    doc.save(path)
    try:
        import docx as _d
        return getattr(_d, "__version__", "python-docx")
    except Exception:
        return "python-docx"


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--audit", default=None, help="path to a specific audit CSV")
    ap.add_argument("--outdir", default=None, help="override output directory")
    args = ap.parse_args()

    audit_path = find_audit(args.audit)
    outdir = Path(args.outdir) if args.outdir else OUT_DIR
    outdir.mkdir(parents=True, exist_ok=True)

    audit = pd.read_csv(audit_path)
    summary = layer_summary(audit)
    tidy = tidy_frame(audit)

    csv_path = outdir / f"appendix_interactional_lexicon_{RUN_DATE}.csv"
    md_path  = outdir / f"appendix_interactional_lexicon_{RUN_DATE}.md"
    docx_path = outdir / f"appendix_interactional_lexicon_{RUN_DATE}.docx"

    tidy.to_csv(csv_path, index=False)
    write_markdown(md_path, summary, tidy)
    docx_ver = write_docx(docx_path, summary, tidy)

    print(f"  audit read  <- {audit_path}")
    print(f"  appendix csv -> {csv_path}")
    print(f"  appendix md  -> {md_path}")
    outs = [str(csv_path), str(md_path)]
    if docx_ver:
        print(f"  appendix docx-> {docx_path}  (python-docx {docx_ver})")
        outs.append(str(docx_path))
    else:
        print("  appendix docx-> SKIPPED (python-docx not installed; "
              "pip install python-docx)")
    print("\n  Summary by layer:")
    print(summary.to_string(index=False))

    if log_run is not None:
        log = log_run(
            __file__,
            inputs=[str(audit_path)],
            outputs=outs,
            params="deterministic; no randomness; formats audit CSV only",
            notes=f"python-docx={docx_ver or 'absent'}; "
                  f"layers={LAYER_ORDER}",
        )
        print(f"  run log -> {log}")


if __name__ == "__main__":
    main()
